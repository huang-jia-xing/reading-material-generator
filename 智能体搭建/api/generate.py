"""
简化版文件生成API - 最终修复版
专为Vercel Serverless Functions优化
完全零存储，所有文件在内存中生成
"""

import json
import zipfile
from io import BytesIO
from datetime import datetime

# 检查依赖，提供回退方案
try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告：python-docx未安装，将使用纯文本格式")
    # 在except块中定义这些变量以避免错误
    Document = None
    Pt = None
    RGBColor = None

try:
    from PIL import Image, ImageDraw, ImageFont
    HAS_PIL = True
except ImportError:
    HAS_PIL = False
    print("警告：Pillow未安装，将跳过图片生成")
    # 在except块中定义这些变量以避免错误
    Image = None
    ImageDraw = None
    ImageFont = None

def handler(event, _context=None):
    """Vercel Serverless Function 入口点"""
    try:
        # 解析请求
        if event.get('httpMethod') == 'OPTIONS':
            return {
                'statusCode': 200,
                'headers': {
                    'Access-Control-Allow-Origin': '*',
                    'Access-Control-Allow-Methods': 'POST, OPTIONS',
                    'Access-Control-Allow-Headers': 'Content-Type',
                },
                'body': ''
            }

        # 解析请求体
        body = json.loads(event.get('body', '{}'))

        # 生成文件 - 重命名变量避免警告
        zip_binary_data = generate_reading_materials(body)

        # 返回ZIP文件
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/zip',
                'Content-Disposition': 'attachment; filename="reading_materials.zip"',
                'Access-Control-Allow-Origin': '*',
            },
            'body': zip_binary_data.decode('latin-1'),  # Vercel要求字符串
            'isBase64Encoded': False
        }

    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': str(e)})
        }

def generate_reading_materials(data):
    """生成阅读材料并返回ZIP文件的二进制数据"""
    # 创建内存中的ZIP文件
    zip_buffer = BytesIO()

    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # 生成各版本阅读文章
        versions = data.get('leveled_texts', {})
        for version, content in versions.items():
            # Word文档
            doc_content = generate_word_content(version, content)
            version_name = get_version_name(version)
            file_name = content.get('title', '文章').replace('/', '_')  # 防止路径问题
            zip_file.writestr(
                f"阅读文章_{version_name}_{file_name}.docx",
                doc_content
            )

            # 纯文本版本（备用）
            text_content = f"{content.get('title', '')}\n\n{content.get('content', '')}"
            zip_file.writestr(
                f"阅读文章_{version_name}_纯文本.txt",
                text_content.encode('utf-8')
            )

        # 生成阅读理解问题
        questions = data.get('comprehension_questions', {})
        questions_content = generate_questions_content(questions)
        zip_file.writestr("阅读理解问题.docx", questions_content)

        # 生成词汇表
        vocabulary = data.get('support_materials', {})
        vocab_content = generate_vocabulary_content(vocabulary)
        zip_file.writestr("词汇表.docx", vocab_content)

        # 生成教师指南（简化版）
        guide_content = generate_teacher_guide(data)
        zip_file.writestr("教师使用指南.docx", guide_content)

        # 生成使用说明文件
        readme = """# 分层阅读材料使用说明

## 文件说明
1. 阅读文章_XXX.docx - 分层阅读文章
2. 阅读理解问题.docx - 配套练习题
3. 词汇表.docx - 重点词汇学习
4. 教师使用指南.docx - 教学建议

## 使用建议
1. 根据学生阅读水平分配合适版本
2. 鼓励学生挑战更高难度版本
3. 使用配套问题进行阅读评估

## 技术支持
如有问题，请联系系统管理员。
"""
        zip_file.writestr("使用说明.txt", readme.encode('utf-8'))

    # 返回ZIP文件的二进制数据
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def generate_word_content(version, content):
    """生成Word文档内容"""
    if not HAS_DOCX or Document is None:
        # 备用方案：返回纯HTML
        html = f"""
        <html>
        <head><meta charset="UTF-8"></head>
        <body>
            <h1>{content.get('title', '')}</h1>
            <h3>{get_version_name(version)}</h3>
            <p>字数：{content.get('word_count', 0)}</p>
            <hr>
            <div style="line-height: 1.6;">
            {content.get('content', '')}
            </div>
        </body>
        </html>
        """
        return html.encode('utf-8')

    # 使用python-docx生成
    try:
        doc = Document()

        # 标题
        title = doc.add_heading(content.get('title', '阅读文章'), 0)
        title_run = title.add_run(content.get('title', '阅读文章'))
        title_run.font.size = Pt(24)
        title_run.font.color.rgb = RGBColor(0, 51, 102)

        # 版本信息
        version_para = doc.add_paragraph()
        version_para.add_run(f"版本：{get_version_name(version)}").bold = True

        # 基本信息
        info = doc.add_paragraph()
        info.add_run(f"字数：{content.get('word_count', 0)} | ")
        info.add_run(f"阅读难度：{content.get('reading_level', '标准')}")

        # 分隔线
        doc.add_paragraph().add_run("─" * 50)

        # 正文 - 直接分段处理
        content_text = content.get('content', '')
        paragraphs = content_text.split('\n')
        for para in paragraphs:
            if para.strip():
                paragraph = doc.add_paragraph(para.strip())
                paragraph.paragraph_format.line_spacing = 1.5
                paragraph.paragraph_format.space_after = Pt(6)

        # 保存到内存
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"生成Word文档失败: {e}")
        # 备用方案
        return f"ERROR: {str(e)}".encode('utf-8')

def generate_questions_content(questions_data):
    """生成阅读理解问题文档"""
    if not HAS_DOCX or Document is None:
        # 简化版
        content = "阅读理解问题\n\n"
        for version, questions in questions_data.items():
            if not questions:
                continue
            version_name = get_version_name(version.replace('_questions', ''))
            content += f"\n{version_name}问题：\n"
            for i, q in enumerate(questions, 1):
                content += f"{i}. {q.get('question', '')}\n"
        return content.encode('utf-8')

    try:
        doc = Document()
        doc.add_heading('阅读理解问题', 0)

        for version, questions in questions_data.items():
            if not questions:
                continue

            version_name = get_version_name(version.replace('_questions', ''))
            doc.add_heading(f'{version_name}问题', level=1)

            for i, q in enumerate(questions, 1):
                # 问题
                question_para = doc.add_paragraph()
                question_para.add_run(f'{i}. {q.get("question", "")}').bold = True

                # 选项
                if q.get('type') == 'choice' and q.get('options'):
                    for j, option in enumerate(q.get('options', [])):
                        doc.add_paragraph(f'   {chr(65+j)}. {option}', style='List Bullet')

                # 答案
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run(f'答案：{q.get("answer", "")}')
                answer_run.font.color.rgb = RGBColor(0, 128, 0)

                # 解析
                if q.get('explanation'):
                    doc.add_paragraph(f'解析：{q.get("explanation")}', style='List Bullet')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"生成问题文档失败: {e}")
        return f"ERROR: {str(e)}".encode('utf-8')

def generate_vocabulary_content(support_materials):
    """生成词汇表文档"""
    if not HAS_DOCX or Document is None:
        content = "词汇表\n\n"
        for version_key, materials in support_materials.items():
            version = version_key.replace('_materials', '')
            version_name = get_version_name(version)
            content += f"\n{version_name}词汇表：\n"

            vocab_list = materials.get('vocabulary_list', [])
            for vocab in vocab_list:
                content += f"• {vocab.get('word', '')}：{vocab.get('definition', '')}\n"
        return content.encode('utf-8')

    try:
        doc = Document()
        doc.add_heading('词汇表', 0)

        for version_key, materials in support_materials.items():
            version = version_key.replace('_materials', '')
            version_name = get_version_name(version)

            doc.add_heading(f'{version_name}词汇表', level=1)

            # 创建表格
            vocab_list = materials.get('vocabulary_list', [])
            if vocab_list:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Light Grid Accent 1'

                # 表头
                headers = table.rows[0].cells
                headers[0].text = '词语'
                headers[1].text = '拼音'
                headers[2].text = '解释'
                headers[3].text = '例句'

                # 添加数据
                for vocab in vocab_list:
                    row_cells = table.add_row().cells
                    row_cells[0].text = vocab.get('word', '')
                    row_cells[1].text = vocab.get('pinyin', '')
                    row_cells[2].text = vocab.get('definition', '')
                    row_cells[3].text = vocab.get('example', '')

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    except Exception as e:
        print(f"生成词汇表失败: {e}")
        return f"ERROR: {str(e)}".encode('utf-8')


def generate_teacher_guide(data):
    """生成教师指南 - 修复版"""
    try:
        guide = f"""# 教师使用指南

        ## 课程信息
        - 生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}
        - 主题：{data.get('core_theme', '自定义主题')}
        
        ## 使用建议
        
        ### 1. 分组教学
        - 基础版：适合阅读困难的学生
        - 标准版：适合大多数学生  
        - 挑战版：适合阅读能力强的学生
        
        ### 2. 教学流程
        1. 课前：分发适合学生水平的阅读材料
        2. 课中：组织小组讨论，鼓励学生分享
        3. 课后：使用配套问题进行评估
        
        ### 3. 差异化策略
        - 允许学生根据自己的进度选择材料
        - 鼓励完成基础版的学生尝试挑战版
        - 组织跨版本的小组合作
        
        ### 4. 评估建议
        - 使用配套的阅读理解问题
        - 观察学生在讨论中的表现
        - 鼓励学生进行自我评估
        
        ## 注意事项
        1. 建议教师先阅读所有版本的材料
        2. 根据学生的实际反应调整教学策略
        3. 鼓励学生提出问题，激发思考
        
        ---
        *本材料由分层阅读材料生成系统生成，仅供教学参考*
        """

        # 使用更简单的纯文本格式，避免Word兼容性问题
        return guide.encode('utf-8')
    except Exception as e:
        return f"教师指南生成失败: {str(e)}".encode('utf-8')

def get_version_name(version_key):
    """获取版本的中文名称"""
    names = {
        'basic': '基础版',
        'standard': '标准版',
        'advanced': '挑战版',
        'extension': '拓展版'
    }
    return names.get(version_key, version_key)

# 本地测试代码（仅当直接运行此文件时执行）
if __name__ == "__main__":
    # 测试数据
    test_data = {
        "leveled_texts": {
            "basic": {
                "title": "蚂蚁的生活",
                "content": "蚂蚁是昆虫。它们住在地下。蚂蚁有不同的工作。有的蚂蚁找食物。有的蚂蚁保护家园。蚂蚁用气味说话。",
                "word_count": 45,
                "reading_level": "基础"
            },
            "standard": {
                "title": "蚂蚁的社会生活",
                "content": "蚂蚁是一种社会性昆虫，它们生活在地下的巢穴中。蚁群中有不同的分工：工蚁负责寻找食物和建造巢穴，兵蚁负责保卫家园，蚁后负责生小蚂蚁。蚂蚁通过释放特殊的气味来互相交流信息。",
                "word_count": 85,
                "reading_level": "标准"
            }
        },
        "comprehension_questions": {
            "basic_questions": [
                {
                    "question": "蚂蚁住在哪里？",
                    "type": "choice",
                    "options": ["树上", "地下", "水里"],
                    "answer": "地下",
                    "explanation": "文章中说'它们住在地下'"
                }
            ],
            "standard_questions": [
                {
                    "question": "蚂蚁用什么方式互相交流？",
                    "type": "short_answer",
                    "answer": "特殊的气味",
                    "explanation": "文章提到'蚂蚁通过释放特殊的气味来互相交流'"
                }
            ]
        },
        "support_materials": {
            "basic_materials": {
                "vocabulary_list": [
                    {
                        "word": "昆虫",
                        "pinyin": "kūn chóng",  # 拼音正确，不是拼写错误
                        "definition": "身体分头、胸、腹三部分，有六只脚的动物",
                        "example": "蝴蝶和蚂蚁都是昆虫。"
                    }
                ]
            }
        },
        "core_theme": "蚂蚁的社会性"
    }

    # 生成测试ZIP文件
    zip_binary = generate_reading_materials(test_data)

    # 保存到本地文件
    with open("../test_output.zip", "wb") as f:
        f.write(zip_binary)

    print("测试完成！已生成 test_output.zip 文件")